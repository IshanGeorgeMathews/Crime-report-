import app.core.paths  # Configures Python path for importing existing modules
import os
import re
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy import delete

from app.infrastructure.documents.utils import load_profile_database, PersonProfile, create_new_profile, save_uo_note, generate_uo_note_text
from app.config import settings
from app.infrastructure.database.models import Profile, ProfileCase, ProfileRelation, ProfileActivity
from app.infrastructure.qdrant.qdrant_service import QdrantService

class ProfileService:
    def __init__(self):
        self.qdrant = QdrantService()

    def _normalize_profile_name(self, value: str) -> str:
        """Return a normalized key for matching profile names to docx filenames."""
        safe_value = re.sub(r'[<>:"/\\|?*]', '_', value or "")
        return safe_value.strip().lower()

    def _extract_name_key_from_filename(self, filename: str) -> Optional[tuple]:
        """Parse a profile docx filename into a normalized name key and review flag."""
        lower = filename.lower()
        if not lower.endswith(".docx"):
            return None
        if "uo " in lower or "uo_" in lower:
            return None
        if "template" in lower or "form details" in lower:
            return None

        stem = filename[:-5]
        is_review = lower.endswith("_review.docx")
        if is_review:
            name_part = stem[:-7]
        else:
            match = re.match(r'^\d+\)\s*(.*)', stem)
            name_part = match.group(1) if match else stem

        name_key = self._normalize_profile_name(name_part)
        if not name_key:
            return None
        return name_key, is_review

    def _resolve_profile_docx_path(self, profile_name: str, prefer_review: bool = False) -> Optional[str]:
        """Find the best matching dossier path for a profile name."""
        target_key = self._normalize_profile_name(profile_name)
        if not target_key or not os.path.isdir(settings.PP_DIR):
            return None

        preferred_match = None
        fallback_match = None
        for filename in os.listdir(settings.PP_DIR):
            parsed = self._extract_name_key_from_filename(filename)
            if not parsed:
                continue

            file_name_key, is_review = parsed
            if file_name_key != target_key:
                continue

            full_path = os.path.join(settings.PP_DIR, filename)
            if is_review == prefer_review and preferred_match is None:
                preferred_match = full_path
            elif fallback_match is None:
                fallback_match = full_path

        return preferred_match or fallback_match

    async def sync_all_profiles_to_db(self, db: AsyncSession):
        """Load all profiles from PP_DIR docx files and sync them to SQL."""
        fs_profiles = load_profile_database(settings.PP_DIR)

        selected_profiles = {}
        for fs_prof in fs_profiles:
            name = fs_prof.name
            if not name:
                name_part = fs_prof.filename[:-5]
                match = re.match(r"^\d+\)\s*(.*)", name_part)
                if match:
                    name_part = match.group(1)
                name = name_part.replace("_review", "").strip().title()

            profile_key = self._normalize_profile_name(name)
            is_review = "_review" in fs_prof.filename.lower()
            existing = selected_profiles.get(profile_key)

            # If both review and production files exist for the same person, keep production.
            if existing and not existing["is_review"]:
                continue
            if existing and existing["is_review"] and is_review:
                continue

            selected_profiles[profile_key] = {
                "profile": fs_prof,
                "resolved_name": name,
                "is_review": is_review,
            }

        for selected in selected_profiles.values():
            fs_prof = selected["profile"]
            # Reconcile status based on filename or registry
            review_status = "approved"
            if "_review" in fs_prof.filename.lower():
                review_status = "pending"
                
            # Parse simple fields from docx
            name = fs_prof.name or selected["resolved_name"]

            # Query if profile already exists in DB by filename or name
            result = await db.execute(select(Profile).filter(Profile.name == name))
            db_prof = result.scalars().first()
            
            # Map docx keys to DB columns
            pp_id_val = fs_prof.pp_id.strip() if fs_prof.pp_id and fs_prof.pp_id.strip() else None
            profile_data = {
                "pp_id": pp_id_val,
                "name": name,
                "parentage": fs_prof.parentage,
                "address": fs_prof.address,
                "police_station": fs_prof.police_station,
                "dob": fs_prof.data.get("DOB", "").strip(),
                "place_of_birth": fs_prof.data.get("Place of Birth", "").strip(),
                "qualification": fs_prof.data.get("Qualification", "").strip(),
                "religion": fs_prof.data.get("Religion", "").strip(),
                "identification_marks": fs_prof.data.get("Identification marks", "").strip(),
                "mobile": fs_prof.data.get("Mobile No", "").strip(),
                "activity_type": fs_prof.activity_type,
                "reason_for_inclusion": fs_prof.data.get("Reason for Inclusion", "").strip(),
                "organization_name": fs_prof.data.get("Organization name", "").strip(),
                "organization_remarks": fs_prof.data.get("Remarks", "").strip(),
                "brief_history": fs_prof.data.get("Brief History of Person", "").strip(),
                "review_status": review_status,
                "neo4j_node_id": f"ind_{name.lower().replace(' ', '_')}"
            }
            
            if not db_prof:
                db_prof = Profile(id=str(uuid.uuid4()), **profile_data)
                db.add(db_prof)
                await db.flush()
            else:
                for k, v in profile_data.items():
                    setattr(db_prof, k, v)
                # Always stamp updated_at so frontend sees fresh data
                db_prof.updated_at = datetime.utcnow()
            
            # Recreate relations in DB
            await db.execute(delete(ProfileRelation).where(ProfileRelation.profile_id == db_prof.id))
            for rel in fs_prof.relations:
                db_rel = ProfileRelation(
                    id=str(uuid.uuid4()),
                    profile_id=db_prof.id,
                    name=rel.get("name", ""),
                    relation_type=rel.get("relation", ""),
                    address=rel.get("address", ""),
                    mobile=rel.get("mobile", "")
                )
                db.add(db_rel)
                
            # Recreate cases in DB
            await db.execute(delete(ProfileCase).where(ProfileCase.profile_id == db_prof.id))
            for case in fs_prof.cases:
                db_case = ProfileCase(
                    id=str(uuid.uuid4()),
                    profile_id=db_prof.id,
                    fir_number=case.get("fir", ""),
                    under_sections=case.get("sections", ""),
                    police_station=case.get("police_station", ""),
                    case_brief=case.get("brief", ""),
                    case_status=case.get("status", "Under Investigation"),
                    co_accused=case.get("co_accused", ""),
                    neo4j_case_node_id=f"case_{case.get('fir', '').lower().replace(' ', '_').replace('/', '_')}"
                )
                db.add(db_case)
                
            # Parse and recreate activities
            await db.execute(delete(ProfileActivity).where(ProfileActivity.profile_id == db_prof.id))
            activity_names = [n.strip() for n in fs_prof.data.get("Name of activity", "").split(";") if n.strip()]
            activity_descs = [d.strip() for d in fs_prof.data.get("Activity Description", "").split(";") if d.strip()]
            activity_dates = [d.strip() for d in fs_prof.data.get("Date of Occurrence", "").split(";") if d.strip()]
            
            # Zip activities together
            num_activities = max(len(activity_names), len(activity_descs), len(activity_dates))
            for i in range(num_activities):
                act_name = activity_names[i] if i < len(activity_names) else "Mentioned in IS Report"
                act_desc = activity_descs[i] if i < len(activity_descs) else ""
                act_date = activity_dates[i] if i < len(activity_dates) else ""
                
                db_act = ProfileActivity(
                    id=str(uuid.uuid4()),
                    profile_id=db_prof.id,
                    activity_name=act_name,
                    activity_desc=act_desc,
                    activity_date=act_date
                )
                db.add(db_act)
            
            # Index in Qdrant
            summary_text = f"{name}. {fs_prof.activity_type}. {fs_prof.address}. PS: {fs_prof.police_station}. History: {profile_data['brief_history']}"
            import asyncio
            await asyncio.to_thread(
                self.qdrant.upsert_item,
                collection="profiles",
                point_id=db_prof.id,
                text=summary_text,
                payload={
                    "profile_id": db_prof.id,
                    "pp_id": fs_prof.pp_id,
                    "name": name,
                    "activity_type": fs_prof.activity_type,
                    "police_station": fs_prof.police_station,
                    "review_status": review_status
                }
            )
            
        await db.commit()

    async def get_profiles(self, db: AsyncSession, page: int = 1, limit: int = 10, search: str = None) -> Dict[str, Any]:
        """Fetch paginated lists of suspect profiles."""
        query = select(Profile)
        
        # Simple search filter if specified
        if search:
            query = query.filter(Profile.name.contains(search) | Profile.pp_id.contains(search) | Profile.activity_type.contains(search))
            
        # Get total count
        result_count = await db.execute(query)
        total = len(result_count.scalars().all())
        
        # Pagination
        query = query.offset((page - 1) * limit).limit(limit)
        result = await db.execute(query)
        items = result.scalars().all()
        
        pages = (total + limit - 1) // limit if total > 0 else 1
        
        return {
            "items": items,
            "total": total,
            "page": page,
            "limit": limit,
            "pages": pages
        }

    async def get_profile_detail(self, db: AsyncSession, profile_id: str) -> Optional[Dict[str, Any]]:
        """Fetch full details of a suspect profile, including cases, relations, and activities."""
        result = await db.execute(select(Profile).filter(Profile.id == profile_id))
        profile = result.scalars().first()
        if not profile:
            return None
            
        # Load relationships
        res_cases = await db.execute(select(ProfileCase).filter(ProfileCase.profile_id == profile_id))
        cases = res_cases.scalars().all()
        
        res_rels = await db.execute(select(ProfileRelation).filter(ProfileRelation.profile_id == profile_id))
        relations = res_rels.scalars().all()
        
        res_acts = await db.execute(select(ProfileActivity).filter(ProfileActivity.profile_id == profile_id))
        activities = res_acts.scalars().all()
        
        # Map profile object to dictionary
        p_dict = {
            "id": profile.id,
            "ppId": profile.pp_id,
            "name": profile.name,
            "parentage": profile.parentage,
            "address": profile.address,
            "policeStation": profile.police_station,
            "dob": profile.dob,
            "placeOfBirth": profile.place_of_birth,
            "qualification": profile.qualification,
            "religion": profile.religion,
            "identificationMarks": profile.identification_marks,
            "mobile": profile.mobile,
            "activityType": profile.activity_type,
            "reasonForInclusion": profile.reason_for_inclusion,
            "organizationName": profile.organization_name,
            "organizationRemarks": profile.organization_remarks,
            "briefHistory": profile.brief_history,
            "reviewStatus": profile.review_status,
            "createdAt": profile.created_at.isoformat() + "Z",
            "updatedAt": profile.updated_at.isoformat() + "Z",
            "cases": cases,
            "relations": relations,
            "activities": activities
        }
        return p_dict

    async def update_profile(self, db: AsyncSession, profile_id: str, updates: Dict[str, Any]) -> Optional[Profile]:
        """Update suspect profile fields in DB, Qdrant, and in the docx on disk."""
        result = await db.execute(select(Profile).filter(Profile.id == profile_id))
        profile = result.scalars().first()
        if not profile:
            return None
            
        # Update DB fields
        field_mappings = {
            "ppId": "pp_id",
            "name": "name",
            "parentage": "parentage",
            "address": "address",
            "policeStation": "police_station",
            "dob": "dob",
            "placeOfBirth": "place_of_birth",
            "qualification": "qualification",
            "religion": "religion",
            "identificationMarks": "identification_marks",
            "mobile": "mobile",
            "activityType": "activity_type",
            "reasonForInclusion": "reason_for_inclusion",
            "organizationName": "organization_name",
            "organizationRemarks": "organization_remarks",
            "briefHistory": "brief_history",
            "reviewStatus": "review_status"
        }
        
        for json_key, db_col in field_mappings.items():
            if json_key in updates:
                val = updates[json_key]
                if db_col == "pp_id" and isinstance(val, str):
                    val = val.strip() if val.strip() else None
                setattr(profile, db_col, val)
                
        profile.updated_at = datetime.utcnow()
        await db.commit()
        
        # Sync back to filesystem docx file
        await self._sync_db_profile_to_docx(db, profile)
        
        # Update Qdrant
        summary_text = f"{profile.name}. {profile.activity_type}. {profile.address}. PS: {profile.police_station}. History: {profile.brief_history}"
        import asyncio
        await asyncio.to_thread(
            self.qdrant.upsert_item,
            collection="profiles",
            point_id=profile.id,
            text=summary_text,
            payload={
                "profile_id": profile.id,
                "pp_id": profile.pp_id,
                "name": profile.name,
                "activity_type": profile.activity_type,
                "police_station": profile.police_station,
                "review_status": profile.review_status
            }
        )
        
        return profile

    async def _sync_db_profile_to_docx(self, db: AsyncSession, profile: Profile):
        """Update the physical docx file on disk with the SQL database values."""
        # Find docx file in PP_DIR
        # Dossier naming convention: "<num>)  <Name>.docx" or "<Name>_review.docx"
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', profile.name)
        target_file = self._resolve_profile_docx_path(
            profile.name,
            prefer_review=profile.review_status == "pending",
        )
                    
        if not target_file:
            # Create a new dossier if missing on disk
            # Find next sequence number
            existing_nums = [
                int(m.group(1))
                for file in os.listdir(settings.PP_DIR)
                for m in [re.match(r'^(\d+)\)', file)]
                if m
            ]
            next_num = max(existing_nums, default=0) + 1
            filename = f"{next_num})  {safe_name}.docx" if profile.review_status == "approved" else f"{safe_name}_review.docx"
            target_file = os.path.join(settings.PP_DIR, filename)
            
            # Pre-create via utility template
            create_new_profile(
                name=profile.name,
                template_path=settings.PP_TEMPLATE,
                output_path=target_file
            )

        try:
            from docx import Document
            doc = Document(target_file)
            
            # Simple field maps inside the paragraphs
            field_map = {
                "PP ID": profile.pp_id or "",
                "Name of Person": profile.name or "",
                "Parentage Name": profile.parentage or "",
                "Address": profile.address or "",
                "Type of Activity": profile.activity_type or "",
                "Reason for Inclusion": profile.reason_for_inclusion or "",
                "Police Station": profile.police_station or "",
                "DOB": profile.dob or "",
                "Place of Birth": profile.place_of_birth or "",
                "Qualification": profile.qualification or "",
                "Religion": profile.religion or "",
                "Identification marks": profile.identification_marks or "",
                "Mobile No": profile.mobile or "",
                "Organization name": profile.organization_name or "",
                "Remarks": profile.organization_remarks or "",
                "Brief History of Person": profile.brief_history or "",
            }
            
            for para in doc.paragraphs:
                txt = para.text.strip()
                for key, val in field_map.items():
                    if txt.startswith(key):
                        para.text = f"{key}\t\t-\t{val}"
                        
            # Sync Activities back to file paragraphs
            # Load activities from DB
            res_acts = await db.execute(select(ProfileActivity).filter(ProfileActivity.profile_id == profile.id))
            acts = res_acts.scalars().all()
            
            act_names = "; ".join([a.activity_name for a in acts])
            act_descs = "; ".join([a.activity_desc for a in acts])
            act_dates = "; ".join([a.activity_date for a in acts])
            
            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt.startswith("Name of activity"):
                    para.text = f"Name of activity\t\t-{act_names}"
                elif txt.startswith("Activity Description"):
                    para.text = f"Activity Description\t-{act_descs}"
                elif txt.startswith("Date of Occurrence"):
                    para.text = f"Date of Occurrence\t-{act_dates}"

            # Sync Tables (Relations & Case Details)
            res_rels = await db.execute(select(ProfileRelation).filter(ProfileRelation.profile_id == profile.id))
            rels = res_rels.scalars().all()
            
            if len(doc.tables) > 0:
                t0 = doc.tables[0]
                # Keep header, clear other rows
                while len(t0.rows) > 1:
                    t0.rows[-1]._element.getparent().remove(t0.rows[-1]._element)
                for r in rels:
                    row = t0.add_row()
                    row.cells[1].text = r.name
                    row.cells[2].text = r.relation_type
                    row.cells[3].text = r.address
                    row.cells[4].text = r.mobile
                    
            res_cases = await db.execute(select(ProfileCase).filter(ProfileCase.profile_id == profile.id))
            cases = res_cases.scalars().all()
            
            if len(doc.tables) > 1:
                t1 = doc.tables[1]
                # Keep header, clear other rows
                while len(t1.rows) > 1:
                    t1.rows[-1]._element.getparent().remove(t1.rows[-1]._element)
                for c in cases:
                    row = t1.add_row()
                    row.cells[1].text = c.fir_number
                    row.cells[2].text = c.under_sections
                    row.cells[3].text = c.police_station
                    row.cells[4].text = c.case_brief
                    row.cells[5].text = c.case_status
                    row.cells[6].text = c.co_accused
                    
            doc.save(target_file)
        except Exception as e:
            print(f"[Warning] Failed to save DB updates to docx file: {e}")

    async def get_profile_docx_path(self, profile_id: str, db: AsyncSession) -> Optional[str]:
        """Get the absolute filepath of the dossier docx file."""
        result = await db.execute(select(Profile).filter(Profile.id == profile_id))
        profile = result.scalars().first()
        if not profile:
            return None

        return self._resolve_profile_docx_path(
            profile.name,
            prefer_review=profile.review_status == "pending",
        )

    async def generate_uo_note(self, profile_id: str, db: AsyncSession) -> Optional[str]:
        """Generate a Malayalam narrative UO Note .docx on-demand and return its path."""
        result = await db.execute(select(Profile).filter(Profile.id == profile_id))
        profile = result.scalars().first()
        if not profile:
            return None
            
        # Read the profile from disk as PersonProfile
        docx_path = await self.get_profile_docx_path(profile_id, db)
        if not docx_path:
            return None
            
        prof_obj = PersonProfile(docx_path)
        
        # Generate note text in Malayalam
        note_text = generate_uo_note_text(
            profile=prof_obj,
            use_ollama=True,
            ollama_url=settings.OLLAMA_URL
        )
        
        # Save note
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', profile.name)
        uo_filename = f"UO_{safe_name}.docx"
        uo_path = os.path.join(settings.PP_DIR, uo_filename)
        
        save_uo_note(note_text, uo_path)
        return uo_path

"""
utils.py - Core utilities for Intelligence Report Consolidation & Profile Sync
===============================================================================
Handles:
  - .docx reading/writing via python-docx
  - Malayalam <-> English translation (IndicTrans2 first, Google fallback)
  - Categorisation of BACK FILES into Events, Forecasts, Social Media
  - Profile database loading, matching, disambiguation, creation & update
  - Daily Report document generation with proper formatting
"""

import os
import re
import json
import unicodedata
import difflib
from datetime import datetime, timedelta
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DISTRICT_CODES = [
    "TVM", "KLM", "PTA", "ALP", "KTM", "IDK", "EKM",
    "TSR", "PKD", "MPM", "KKD", "WYD", "KNR", "KSD",
]

SOCIAL_MEDIA_KEYWORDS = ["socialmedia", "social media", "social_media"]

# Malayalam Unicode range: U+0D00 – U+0D7F
_ML_RE = re.compile(r"[\u0D00-\u0D7F]")

HONORIFICS = {"sri", "sri.", "shri", "shri.", "smt", "smt.", "mr", "mr.", "mrs", "mrs.", "ms", "ms.", "dr", "dr."}

# ---------------------------------------------------------------------------
# Legacy stubs – kept empty so downstream imports (graph_db, cleanup) don't
# break.  Classification is now 100% model-driven via NEREngine.
# ---------------------------------------------------------------------------
PLACE_NAMES: set = set()   # DEPRECATED – do not add entries
JUNK_WORDS:  set = set()   # DEPRECATED – do not add entries
STOP_WORDS:  set = set()   # DEPRECATED – alias


def _is_valid_person_name(name: str) -> bool:
    """Structural heuristic: rejects obviously non-name strings.

    This does NOT do vocabulary / dictionary look-ups.  Semantic
    classification (person vs. place vs. junk) is handled exclusively
    by the LLM in NEREngine.classify_candidates().
    """
    if not name or len(name.strip()) < 3:
        return False

    tokens = [t.strip(".,()[]{}'\"-:") for t in name.split() if t.strip(".,()[]{}'\"-:")]
    if not tokens:
        return False

    # Reject pure acronyms (e.g. RSU, KMP, UAPA)
    for tok in tokens:
        if tok.isupper() and 2 <= len(tok) <= 5:
            return False

    # Reject if contains digits or special characters (not a name)
    joined = " ".join(tokens)
    if re.search(r'\d', joined):
        return False

    # Single-word names must be >= 4 chars
    if len(tokens) == 1 and len(tokens[0]) < 4:
        return False

    # Reject if more than 5 words (sentence fragment, not a name)
    if len(tokens) > 5:
        return False

    return True


# ---------------------------------------------------------------------------
# Language helpers
# ---------------------------------------------------------------------------

def is_malayalam(text: str) -> bool:
    """Return True if more than 30% of alphabetic characters are Malayalam."""
    if not text:
        return False
    ml_count = len(_ML_RE.findall(text))
    alpha_count = sum(1 for c in text if unicodedata.category(c).startswith("L"))
    if alpha_count == 0:
        return False
    return (ml_count / alpha_count) > 0.30


def sanitize_path(path: str) -> str:
    """Check for path traversal and resolve absolute path."""
    if not path:
        return path
    if ".." in path:
        raise ValueError(f"Path traversal detected in path: {path}")
    return os.path.abspath(path)


def filter_control_characters(text: str) -> str:
    """Remove control characters from text, preserving tab, newline, carriage return."""
    if not text:
        return text
    return "".join(c for c in text if not (unicodedata.category(c) == "Cc" and c not in "\n\t\r"))


def check_file_size(path: str, max_mb: float = 10.0):
    """Raise ValueError if the file size exceeds limit."""
    if not os.path.exists(path):
        return
    size_mb = os.path.getsize(path) / (1024.0 * 1024.0)
    if size_mb > max_mb:
        raise ValueError(f"File size exceeds limit of {max_mb}MB: {path} ({size_mb:.2f}MB)")


def scan_for_unanchored_suffix_warnings(original_ml: str, translated_en: str) -> list[str]:
    """Scan for cases where Malayalam contains 'കുട്ടി' or 'പിള്ള' but English contains 'child' or 'children'.
    
    This indicates a possible mistranslated suspect name.
    """
    warnings = []
    has_ml = "കുട്ടി" in original_ml or "പിള്ള" in original_ml
    has_en = re.search(r"\bchild(ren)?\b", translated_en, re.IGNORECASE)
    if has_ml and has_en:
        warnings.append(
            f"Actionable Suffix Warning: Detected 'കുട്ടി'/'പിള്ള' in Malayalam source "
            f"translating to 'child'/'children' in English. Verify that a suspect name was not "
            f"mistranslated as a common noun."
        )
    return warnings


def translate_ml_to_en(text: str, use_ollama: bool = False, ollama_url: str = "http://localhost:11434", batch_engines: list = None) -> str:
    """Translate Malayalam text to English using the production TranslationEngine.

    Parameters
    ----------
    text : str
        Source text (may be mixed Malayalam / English).
    use_ollama : bool
        Compatibility flag; Ollama is used by the caller for summarization,
        while raw translation uses IndicTrans2 first.
    ollama_url : str
        Base URL of a running Ollama instance.
    batch_engines : list, optional
        List to track the engines used for translation.

    Returns
    -------
    str
        Translated English text.
    """
    if not text or not text.strip():
        return text

    # If text is already mostly English, return as-is
    if not is_malayalam(text):
        return text

    from translation import TranslationEngine
    engine = TranslationEngine()
    
    # We translate the document which handles script segmentation, suffix anchoring, and chunking
    translated = engine.translate_document(text, use_ollama=use_ollama, ollama_url=ollama_url, batch_engines=batch_engines)
    
    # Run actionable suffix warning scanner
    warnings = scan_for_unanchored_suffix_warnings(text, translated)
    for w in warnings:
        print(f"  [Warning] {w}")
        
    return translated


def _chunk_text(text: str, max_len: int = 4500) -> list:
    """Split text into chunks of at most *max_len* characters at sentence boundaries."""
    if len(text) <= max_len:
        return [text]
    chunks, current = [], ""
    for sentence in re.split(r"(?<=[.!?\u0D7F])\s+", text):
        if len(current) + len(sentence) + 1 > max_len:
            if current:
                chunks.append(current)
            current = sentence
        else:
            current = f"{current} {sentence}".strip()
    if current:
        chunks.append(current)
    return chunks or [text]


# ---------------------------------------------------------------------------
# Docx reading helpers
# ---------------------------------------------------------------------------

def read_docx_full_text(path: str) -> str:
    """Return the full plain-text content of a .docx file with security checks."""
    sanitized = sanitize_path(path)
    check_file_size(sanitized)
    doc = Document(sanitized)
    lines = []
    for para in doc.paragraphs:
        lines.append(filter_control_characters(para.text))
    # Also capture text inside tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = [filter_control_characters(cell.text.strip()) for cell in row.cells]
            lines.append("\t".join(row_texts))
    return "\n".join(lines)


def read_docx_paragraphs(path: str) -> list:
    """Return a list of non-empty paragraph strings from a .docx file with security checks."""
    sanitized = sanitize_path(path)
    check_file_size(sanitized)
    doc = Document(sanitized)
    return [filter_control_characters(p.text.strip()) for p in doc.paragraphs if p.text.strip()]


# ---------------------------------------------------------------------------
# Back-file categorisation
# ---------------------------------------------------------------------------

def categorise_back_files(date_folder: str) -> dict:
    """Categorise files in a BACK FILES/<date> folder.

    Returns
    -------
    dict with keys 'events', 'forecasts', 'social_media', each a list of
    absolute file paths.
    """
    result = {"events": [], "forecasts": [], "social_media": []}

    if not os.path.isdir(date_folder):
        return result

    def natural_sort_key(s):
        return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

    for fname in sorted(os.listdir(date_folder), key=natural_sort_key):
        fpath = os.path.join(date_folder, fname)
        if not os.path.isfile(fpath):
            continue
        if not fname.lower().endswith(".docx"):
            continue

        lower = fname.lower().replace(" ", "")

        # Social media
        if any(kw in lower for kw in SOCIAL_MEDIA_KEYWORDS):
            result["social_media"].append(fpath)
            continue

        # Forecast: district-code files or F-prefixed files
        base_no_ext = os.path.splitext(fname)[0].strip()
        if base_no_ext.upper() in DISTRICT_CODES:
            result["forecasts"].append(fpath)
            continue
        # F-prefixed  (F1, F2, F10, etc.)  or  named forecast location files
        if re.match(r"^[Ff]\d", base_no_ext):
            result["forecasts"].append(fpath)
            continue
        # Named location forecasts (all-caps single-word that are NOT district codes)
        # e.g. VELLAYAMBALAM.docx, PARASSALA.docx, ANANDAVALLISWARAM.docx, KILIKOLLOOR.docx
        if base_no_ext.isupper() and len(base_no_ext) > 3 and base_no_ext not in DISTRICT_CODES:
            result["forecasts"].append(fpath)
            continue

        # Content-based fallback check if filename-based fails to categorise specifically
        try:
            # Inline import of Document to read the docx file content
            from docx import Document
            doc = Document(fpath)
            text = "\n".join(p.text for p in doc.paragraphs)
            first_lines = "\n".join(text.split("\n")[:5]).upper()
            
            if "/RSU/" in first_lines:
                result["social_media"].append(fpath)
                continue
            elif "/CC/" in first_lines:
                result["forecasts"].append(fpath)
                continue
            elif "/EC/" in first_lines:
                result["events"].append(fpath)
                continue
            
            # Content keyword searches
            text_lower = text.lower()
            if any(kw in text_lower for kw in ["social media", "facebook", "twitter", "whatsapp", "telegram", "viral", "rsu", "cyber"]):
                result["social_media"].append(fpath)
                continue
            elif any(kw in text_lower for kw in ["forecast", "anticipated", "alert", "programme", "scheduled", "likely", "expected", "future"]):
                result["forecasts"].append(fpath)
                continue
        except Exception as e:
            print(f"  [Warning] Content-based categorization failed for {fname}: {e}")

        # Everything else → events
        result["events"].append(fpath)

    return result


# ---------------------------------------------------------------------------
# Extract district-category tag from text  e.g. (KLM-EC), (TVM RL-CC)
# ---------------------------------------------------------------------------

_TAG_RE = re.compile(r"\(([A-Z]{2,4}[\s\-]?(?:City|RL|C)?[\s\-]?(?:EC|CC|OC)[^\)]*)\)\s*$")

def extract_district_tag(text: str) -> str:
    """Try to extract a district-category tag from the end of a text block."""
    m = _TAG_RE.search(text)
    if m:
        return m.group(0)
    # Also try a looser search
    m2 = re.search(r"\(([A-Z]{2,4}[^)]{0,15})\)\s*$", text)
    if m2:
        return m2.group(0)
    return ""


# ---------------------------------------------------------------------------
# Profile database
# ---------------------------------------------------------------------------

class PersonProfile:
    """In-memory representation of a PP profile .docx."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
        self.data = {}  # key-value pairs extracted from the profile
        self._parse()

    def _parse(self):
        doc = Document(self.filepath)
        kv_re = re.compile(r"^(.+?)\s*[-–]\s*(.*)$")
        for p in doc.paragraphs:
            txt = p.text.strip()
            m = kv_re.match(txt)
            if m:
                key = m.group(1).strip()
                val = m.group(2).strip()
                self.data[key] = val

        self.relations = []
        self.cases = []
        
        # Parse Relations Table (Table 0)
        if len(doc.tables) > 0:
            t0 = doc.tables[0]
            for row in t0.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 5 and any(cells[1:]):
                    self.relations.append({
                        "name": cells[1],
                        "relation": cells[2],
                        "address": cells[3],
                        "mobile": cells[4]
                    })
                    
        # Parse Case Details Table (Table 1)
        if len(doc.tables) > 1:
            t1 = doc.tables[1]
            for row in t1.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 7 and any(cells[1:]):
                    self.cases.append({
                        "fir": cells[1],
                        "sections": cells[2],
                        "police_station": cells[3],
                        "brief": cells[4],
                        "status": cells[5],
                        "co_accused": cells[6]
                    })

    @property
    def name(self) -> str:
        return self.data.get("Name of Person", "").strip()

    @property
    def parentage(self) -> str:
        return self.data.get("Parentage Name", "").strip()

    @property
    def address(self) -> str:
        return self.data.get("Address", "").strip()

    @property
    def police_station(self) -> str:
        return self.data.get("Police Station", "").strip()

    @property
    def pp_id(self) -> str:
        return self.data.get("PP ID", "").strip()

    @property
    def activity_type(self) -> str:
        return self.data.get("Type of Activity", "").strip()


def load_profile_database(pp_dir: str) -> list:
    """Load all PP profiles from the PP & Uo Note Dummy directory.

    Returns a list of PersonProfile objects (skipping UO and template files).
    """
    profiles = []
    if not os.path.isdir(pp_dir):
        return profiles
    for fname in os.listdir(pp_dir):
        lower = fname.lower()
        if not lower.endswith(".docx"):
            continue
        if "uo " in lower or "uo_" in lower:
            continue
        if "form details" in lower or "template" in lower:
            continue
        fpath = os.path.join(pp_dir, fname)
        try:
            profiles.append(PersonProfile(fpath))
        except Exception as e:
            print(f"  [Warning] Could not parse profile {fname}: {e}")
    return profiles


# ---------------------------------------------------------------------------
# Name matching & disambiguation
# ---------------------------------------------------------------------------

def is_fuzzy_match(name1: str, name2: str) -> bool:
    """Return True if name1 and name2 are fuzzy or phonetic matches."""
    n1 = name1.lower().strip()
    n2 = name2.lower().strip()
    if n1 == n2:
        return True
    
    # Substring check only if BOTH names are multi-word to prevent single-word false positives
    if " " in n1 and " " in n2:
        if len(n1) > 4 and len(n2) > 4:
            if n1 in n2 or n2 in n1:
                return True
            
    # Visual sequence similarity
    ratio = difflib.SequenceMatcher(None, n1, n2).ratio()
    if ratio >= 0.80:
        return True
        
    # Malayalam phonetic/anglicisation normalization
    def normalize_phonetics(s):
        s = s.replace("oo", "u").replace("ee", "i").replace("tt", "t").replace("sh", "s")
        s = s.replace("ch", "c").replace("krishnan", "krishnn").replace("jecob", "jacob")
        # De-duplicate letters
        res = []
        for char in s:
            if not res or res[-1] != char:
                res.append(char)
        return "".join(res)
        
    if normalize_phonetics(n1) == normalize_phonetics(n2):
        return True
        
    # Overlap of multiple words (e.g. Beena Rasheed Sasi and Sasi Beena)
    w1 = set(n1.split())
    w2 = set(n2.split())
    if len(w1.intersection(w2)) >= 2:
        diff1 = w1 - w2
        diff2 = w2 - w1
        if diff1 and diff2:
            # Helper to check if two words fuzzy match
            def word_match(wd1, wd2):
                if wd1 == wd2:
                    return True
                if difflib.SequenceMatcher(None, wd1, wd2).ratio() >= 0.80:
                    return True
                def norm(s):
                    s = s.replace("oo", "u").replace("ee", "i").replace("tt", "t").replace("sh", "s")
                    s = s.replace("ch", "c").replace("krishnan", "krishnn").replace("jecob", "jacob")
                    res = []
                    for char in s:
                        if not res or res[-1] != char:
                            res.append(char)
                    return "".join(res)
                return norm(wd1) == norm(wd2)
            
            has_close_diff = False
            for wd1 in diff1:
                for wd2 in diff2:
                    if word_match(wd1, wd2):
                        has_close_diff = True
                        break
                if has_close_diff:
                    break
            if not has_close_diff:
                return False
        return True
        
    return False


def extract_candidate_names_from_text(text: str) -> list:
    """Extract and clean sequences of Title Case words representing candidate names from text.
    
    Key improvements:
    - Clean translation typings (like Sri.S.Bibin -> Sri S Bibin)
    - Split text on punctuation (commas, semicolons, parentheses) and written numbers/prepositions
      *before* extracting Title Case words to prevent merged name lists.
    - Filter candidate sequences using structural heuristics via _is_valid_person_name.
    - Restrict name sequences to at most 3 words.
    """
    if not text:
        return []

    # Clean translation typings (like Sri.S.Bibin -> Sri S Bibin)
    text_clean = re.sub(r'(?<=[a-zA-Z])\.(?=[a-zA-Z])', ' ', text)

    # Split text into clauses/phrases using separators to avoid merging lists of names
    separators = r'[,;()]|\b(?:one|two|three|four|five|six|seven|eight|nine|ten)\b|\b(?:and|or|under|in|from|to|at|w/o|s/o|d/o|for|against|demanding|started|commenced|continuing|following|with|of|during|about|by|after)\b|\b\d+[).]\s*'
    parts = re.split(separators, text_clean, flags=re.IGNORECASE)
    
    candidates = []
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
        
        # Extract consecutive sequences of Title Case words
        words = part.split()
        current_seq = []
        for w in words:
            # Clean punctuation from start/end of the word
            w_clean = w.strip(".,()[]{}'\"-:")
            # A Title Case word starts with uppercase, followed by letters
            if w_clean and w_clean[0].isupper() and w_clean.isalpha():
                current_seq.append(w_clean)
            else:
                if current_seq:
                    candidates.append(" ".join(current_seq))
                    current_seq = []
        if current_seq:
            candidates.append(" ".join(current_seq))
            
    filtered_candidates = []
    for cand in candidates:
        tokens = cand.split()
        
        # Clean honorifics from the beginning of the candidate
        # e.g., "Sri. Chittoor Kutty" -> "Chittoor Kutty"
        clean_tokens = []
        for tok in tokens:
            tok_low = tok.lower().strip(".,")
            if tok_low in HONORIFICS:
                continue
            clean_tokens.append(tok)
            
        if not clean_tokens:
            continue
            
        # Limit length to at most 4 words. If more, it's likely a sentence fragment
        if len(clean_tokens) > 4:
            continue
            
        cleaned_name = " ".join(clean_tokens)
        if _is_valid_person_name(cleaned_name):
            filtered_candidates.append(cleaned_name)
        
    return list(set(filtered_candidates))


def find_matching_profiles(text: str, profiles: list) -> list:
    """Return a list of (profile, score) tuples for profiles whose name appears in *text* using fuzzy logic."""
    results = []
    candidates = extract_candidate_names_from_text(text)
    
    for prof in profiles:
        if not prof.name:
            continue
            
        matched_cand = None
        for cand in candidates:
            if is_fuzzy_match(cand, prof.name):
                matched_cand = cand
                break
                
        if not matched_cand:
            continue
            
        score = 10  # base score
        
        # Exact match bonus
        if matched_cand.lower() == prof.name.lower():
            score += 5
            
        # Parentage match (e.g. S/o Sivadasan)
        parent_matches = re.findall(r"[SsDd]/o\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,2})", text)
        if prof.parentage and parent_matches:
            for pm in parent_matches:
                if is_fuzzy_match(pm, prof.parentage):
                    score += 10
                    break
                    
        # Police station match
        if prof.police_station:
            ps_name = prof.police_station.replace(" PS", "").replace(" ps", "").strip()
            if ps_name.lower() in text.lower():
                score += 5
                
        # Address word match
        if prof.address:
            addr_words = [w for w in prof.address.replace(",", " ").split() if len(w) > 3]
            for w in addr_words:
                if w.lower() in text.lower():
                    score += 2
                    break
                    
        results.append((prof, score))
        
    results.sort(key=lambda x: x[1], reverse=True)
    return results


def disambiguate_match(text: str, profiles: list) -> tuple:
    """Find the best-matching profile for a given text.

    Returns
    -------
    (PersonProfile or None, bool is_ambiguous)
    """
    matches = find_matching_profiles(text, profiles)
    if not matches:
        return None, False
    if len(matches) == 1:
        return matches[0][0], False
    # Multiple matches – check if top score is clearly the best
    if matches[0][1] > matches[1][1]:
        return matches[0][0], False
    # Scores are tied → ambiguous
    return matches[0][0], True


def find_profiles_for_name(name: str, profiles: list) -> list:
    """Find all profiles whose stored name fuzzy-matches a specific extracted name.

    Unlike find_matching_profiles() which scans a whole text for any profile,
    this function checks only whether `name` matches each profile's name field.

    Returns
    -------
    list of PersonProfile objects that match `name`.
    """
    results = []
    for prof in profiles:
        if prof.name and is_fuzzy_match(name, prof.name):
            results.append(prof)
    return results


# ---------------------------------------------------------------------------
# Profile creation & update
# ---------------------------------------------------------------------------

def create_new_profile(
    name: str,
    parentage: str = "",
    address: str = "",
    police_station: str = "",
    activity_type: str = "",
    reason: str = "",
    activity_name: str = "",
    activity_desc: str = "",
    activity_date: str = "",
    case_number: str = "",
    case_sections: str = "",
    case_brief: str = "",
    template_path: str = "",
    output_path: str = "",
) -> str:
    """Create a new PP profile .docx from the template, pre-populating available fields.

    Returns the output file path.
    """
    if template_path and os.path.isfile(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    # We'll rebuild by replacing placeholder values in the paragraphs
    field_map = {
        "PP ID": "",
        "Name of Person": name,
        "Parentage Name": parentage,
        "Address": address,
        "Type of Activity": activity_type,
        "Reason for Inclusion": reason,
        "Police Station": police_station,
        "DOB": "",
        "Place of Birth": "",
        "Qualification": "",
        "Religion": "",
        "Identification marks": "",
        "Mobile No": "",
        "Name of activity": activity_name,
        "Activity Description": activity_desc,
        "Date of Occurrence": activity_date,
        "Persons Involved Details": "",
    }

    for para in doc.paragraphs:
        for field_key, field_val in field_map.items():
            if para.text.strip().startswith(field_key):
                # Preserve key, replace value
                para.text = f"{field_key}-{field_val}"

    doc.save(output_path)
    return output_path


def _append_to_relations_table(table, rel_data: dict):
    """Append a relative entry to Table 0 (Relations), filling empty rows first."""
    name     = _to_str(rel_data.get("name")).strip()
    relation = _to_str(rel_data.get("relation")).strip()
    address  = _to_str(rel_data.get("address")).strip()
    mobile   = _to_str(rel_data.get("mobile")).strip()
    if not name:
        return

    # Check if this relative is already in the table to avoid duplicates
    existing_count = 0
    first_empty_row = None
    
    # Iterate through rows starting from index 1 (skip header)
    for row in table.rows[1:]:
        row_cells = [c.text.strip() for c in row.cells]
        # Check if the row is empty (ignore first cell which is SL No)
        is_empty = all(c == "" for c in row_cells[1:])
        if is_empty:
            if first_empty_row is None:
                first_empty_row = row
        else:
            existing_count += 1
            # Check if this exact relative name already exists
            if row_cells[1].lower() == name.lower():
                return  # Duplicate, skip

    # Select or create the row
    if first_empty_row is not None:
        target_row = first_empty_row
    else:
        target_row = table.add_row()

    # Populate cells
    sl_no = str(existing_count + 1)
    target_row.cells[0].text = sl_no
    target_row.cells[1].text = name
    target_row.cells[2].text = relation
    target_row.cells[3].text = address
    target_row.cells[4].text = mobile


def _to_str(val) -> str:
    """Coerce LLM output to a plain string (handles lists, None, etc.)."""
    if val is None:
        return ""
    if isinstance(val, list):
        return ", ".join(str(v) for v in val)
    return str(val)


def _append_to_case_table(table, case_data: dict):
    """Append a case entry to Table 1 (Case Details), filling empty rows first."""
    fir        = _to_str(case_data.get("fir_number")).strip()
    sections   = _to_str(case_data.get("under_sections")).strip()
    ps         = _to_str(case_data.get("police_station")).strip()
    brief      = _to_str(case_data.get("case_brief")).strip()
    status     = _to_str(case_data.get("case_status") or "Under Investigation").strip()
    co_accused = _to_str(case_data.get("co_accused")).strip()
    
    # We need at least one of these to consider it a valid case entry
    if not (fir or sections or brief):
        return

    existing_count = 0
    first_empty_row = None
    
    for row in table.rows[1:]:
        row_cells = [c.text.strip() for c in row.cells]
        is_empty = all(c == "" for c in row_cells[1:])
        if is_empty:
            if first_empty_row is None:
                first_empty_row = row
        else:
            existing_count += 1
            # Check if this exact case already exists
            if fir and row_cells[1].lower() == fir.lower():
                return  # Duplicate, skip

    if first_empty_row is not None:
        target_row = first_empty_row
    else:
        target_row = table.add_row()

    sl_no = str(existing_count + 1)
    target_row.cells[0].text = sl_no
    target_row.cells[1].text = fir
    target_row.cells[2].text = sections
    target_row.cells[3].text = ps
    target_row.cells[4].text = brief
    target_row.cells[5].text = status
    target_row.cells[6].text = co_accused


def _dedup_semicolon_field(existing: str, new_value: str) -> str:
    """Append new_value to a semicolon-separated field only if it's not already present.

    Parameters
    ----------
    existing : str
        Current field content (may contain multiple "; "-separated entries).
    new_value : str
        The new entry to append.

    Returns
    -------
    str
        The merged field content, with new_value appended only if it wasn't
        already present (exact match after stripping).
    """
    if not new_value:
        return existing
    if not existing:
        return new_value

    entries = [e.strip() for e in existing.split("; ") if e.strip()]
    new_clean = new_value.strip()

    # Exact match check
    if new_clean in entries:
        return existing

    entries.append(new_clean)
    return "; ".join(entries)


def _dedup_description_field(existing: str, new_value: str, match_len: int = 80) -> str:
    """Append new_value to a description field only if no existing entry
    shares its first `match_len` characters (fuzzy dedup).

    The LLM may produce slightly different summaries of the same event
    across runs, but the first ~80 chars are usually stable.  This avoids
    appending the same report paragraph 8 times.

    Parameters
    ----------
    existing : str
        Current field content (may contain multiple "; "-separated entries).
    new_value : str
        The new description entry to append.
    match_len : int
        Number of leading characters to compare for fuzzy matching.

    Returns
    -------
    str
        The merged field content, with new_value appended only if no
        existing entry shares its first match_len characters.
    """
    if not new_value:
        return existing
    if not existing:
        return new_value

    entries = [e.strip() for e in existing.split("; ") if e.strip()]
    new_clean = new_value.strip()
    new_prefix = new_clean[:match_len].lower()

    for entry in entries:
        # Exact match
        if entry == new_clean:
            return existing
        # Fuzzy prefix match — same event, slightly different tail
        if entry[:match_len].lower() == new_prefix:
            return existing

    entries.append(new_clean)
    return "; ".join(entries)


def update_profile_activity(
    profile_path: str,
    activity_name: str = "",
    activity_desc: str = "",
    activity_date: str = "",
    persons_involved: str = "",
    structured_data: dict = None,
):
    """Update the Activities and other structured sections in an existing PP profile .docx."""
    doc = Document(profile_path)

    # 1. Update personal details and other text fields if structured_data is provided
    if structured_data:
        personal = structured_data.get("personal_details", {})
        if isinstance(personal, list):
            personal = personal[0] if personal else {}
        elif not isinstance(personal, dict):
            personal = {}

        org = structured_data.get("organization_involvement", {})
        if isinstance(org, list):
            org = org[0] if org else {}
        elif not isinstance(org, dict):
            org = {}

        brief_hist = _to_str(structured_data.get("brief_history")).strip()

        # Map target fields in the profile paragraphs (safely preserve existing data)
        for para in doc.paragraphs:
            txt = para.text.strip()
            
            # Check for parentage
            if txt.startswith("Parentage Name"):
                existing = txt.split("-", 1)[-1].strip()
                parent_val = _to_str(personal.get("parentage")).strip()
                if (not existing or existing == "-" or existing == "") and parent_val:
                    para.text = f"Parentage Name\t\t-\t{parent_val}"
                    
            # Check for address
            elif txt.startswith("Address"):
                existing = txt.split("-", 1)[-1].strip()
                addr_val = _to_str(personal.get("address")).strip()
                if (not existing or existing == "-" or existing == "") and addr_val:
                    para.text = f"Address\t\t\t-\t{addr_val}"

            # Check for police station
            elif txt.startswith("Police Station"):
                existing = txt.split("-", 1)[-1].strip()
                ps_val = _to_str(personal.get("police_station")).strip()
                if (not existing or existing == "-" or existing == "") and ps_val:
                    para.text = f"Police Station\t\t-\t{ps_val}"

            # Check for organization name
            elif txt.startswith("Organization name"):
                existing = txt.split("-", 1)[-1].strip()
                org_val = _to_str(org.get("org_name")).strip()
                if (not existing or existing == "-" or existing == "") and org_val:
                    para.text = f"Organization name\t-\t{org_val}"

            # Check for organization remarks
            elif txt.startswith("Remarks") and para.text.startswith("Remarks"):
                existing = txt.split("-", 1)[-1].strip()
                rem_val = _to_str(org.get("remarks")).strip()
                if (not existing or existing == "-" or existing == "") and rem_val:
                    para.text = f"Remarks\t\t\t-\t{rem_val}"

            # Check for brief history of person
            elif txt.startswith("Brief History of Person"):
                existing = txt.split("-", 1)[-1].strip()
                if (not existing or existing == "-" or existing == "") and brief_hist:
                    para.text = f"Brief History of Person \t-\t{brief_hist}"

    # 2. Update Activities section paragraphs (with deduplication)
    for para in doc.paragraphs:
        txt = para.text.strip()

        if txt.startswith("Name of activity"):
            existing = txt.split("-", 1)[-1].strip()
            merged = _dedup_semicolon_field(existing, activity_name)
            if merged != existing or (not existing and activity_name):
                para.text = f"Name of activity\t\t-{merged}"

        elif txt.startswith("Activity Description"):
            existing = txt.split("-", 1)[-1].strip()
            merged = _dedup_description_field(existing, activity_desc)
            if merged != existing or (not existing and activity_desc):
                para.text = f"Activity Description\t-{merged}"

        elif txt.startswith("Date of Occurrence"):
            existing = txt.split("-", 1)[-1].strip()
            merged = _dedup_semicolon_field(existing, activity_date)
            if merged != existing or (not existing and activity_date):
                para.text = f"Date of Occurrence\t-{merged}"

        elif txt.startswith("Persons Involved"):
            existing = txt.split("-", 1)[-1].strip()
            merged = _dedup_semicolon_field(existing, persons_involved)
            if merged != existing or (not existing and persons_involved):
                para.text = f"Persons Involved Details\t-{merged}"

    # 3. Update Table 0 (Relations) and Table 1 (Case Details)
    if structured_data:
        relations_list = structured_data.get("relations", [])
        if not isinstance(relations_list, list):
            relations_list = []
            
        case_data = structured_data.get("case_details", {})
        if isinstance(case_data, list):
            case_data = case_data[0] if case_data else {}
        elif not isinstance(case_data, dict):
            case_data = {}
        
        # Table 0 is Relations (usually the first table in doc.tables)
        if len(doc.tables) > 0 and relations_list:
            for rel in relations_list:
                _append_to_relations_table(doc.tables[0], rel)
                
        # Table 1 is Case Details (usually the second table in doc.tables)
        if len(doc.tables) > 1 and case_data:
            _append_to_case_table(doc.tables[1], case_data)

    try:
        doc.save(profile_path)
    except PermissionError:
        print(f"  [Warning] Permission denied when saving profile: {profile_path}. It might be open in another application like Microsoft Word.")
    except Exception as e:
        print(f"  [Warning] Failed to save profile {profile_path}: {e}")


# ---------------------------------------------------------------------------
# UO Note generation
# ---------------------------------------------------------------------------

def generate_uo_note_text(profile: PersonProfile, use_ollama: bool = False, ollama_url: str = "http://localhost:11434") -> str:
    """Generate a Malayalam narrative UO Note from a PersonProfile.

    Uses Ollama if available, else constructs a template-based note.
    """
    # Build a structured English summary to convert
    lines = []
    lines.append(f"Name: {profile.name} (PP ID: {profile.pp_id})")
    if profile.parentage:
        lines.append(f"Parentage: {profile.parentage}")
    if profile.address:
        lines.append(f"Address: {profile.address}")
    if profile.activity_type:
        lines.append(f"Type of Activity: {profile.activity_type}")

    # Include case details
    for key in ["Reason for Inclusion", "Name of activity", "Activity Description",
                 "Date of Occurrence", "Current Doing", "Brief History of Person"]:
        val = profile.data.get(key, "").strip()
        if val:
            lines.append(f"{key}: {val}")

    # Include network relations
    if hasattr(profile, "relations") and profile.relations:
        lines.append("Relations:")
        for rel in profile.relations:
            lines.append(f"  - {rel['name']} ({rel['relation']}), Address: {rel['address']}, Mobile: {rel['mobile']}")

    # Include case/crime details history
    if hasattr(profile, "cases") and profile.cases:
        lines.append("Criminal Case History:")
        for c in profile.cases:
            lines.append(f"  - FIR: {c['fir']}, Sections: {c['sections']}, Station: {c['police_station']}, Brief: {c['brief']}, Status: {c['status']}, Co-accused: {c['co_accused']}")

    summary = "\n".join(lines)

    if use_ollama:
        try:
            import requests
            prompt = (
                "Based on the following structured profile data of a person of interest, "
                "generate an official Malayalam Un-Official Note (അനൗദ്യോഗിക കുറിപ്പ്) "
                "summarising their details in narrative paragraph form. "
                "Use formal administrative Malayalam. Preserve all names, dates, "
                "case numbers, and legal sections exactly as given.\n\n"
                f"{summary}"
            )
            resp = requests.post(
                f"{ollama_url}/api/generate",
                json={"model": "llama3", "prompt": prompt, "stream": False},
                timeout=90,
            )
            if resp.status_code == 200:
                return resp.json().get("response", "").strip()
        except Exception:
            pass

    # Fallback: translate the English summary to Malayalam
    try:
        from deep_translator import GoogleTranslator
        header = "അനൗദ്യോഗിക കുറിപ്പ്\n\n"
        translated = GoogleTranslator(source="en", target="ml").translate(summary)
        return header + (translated or summary)
    except Exception:
        return f"അനൗദ്യോഗിക കുറിപ്പ്\n\n{summary}"


def save_uo_note(text: str, output_path: str):
    """Save a UO Note string to a .docx file.

    Handles mixed Malayalam/English content safely — python-docx stores text
    internally as Unicode so no charmap encoding issues arise.
    """
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Nirmala UI"
    font.size = Pt(11)

    for line in text.split("\n"):
        # Ensure line is a proper str (handle any bytes that slipped through)
        if isinstance(line, bytes):
            line = line.decode("utf-8", errors="replace")
        # Replace Windows-1252 smart-quote surrogates that can't encode
        safe_line = line.encode("utf-8", errors="replace").decode("utf-8", errors="replace")
        para = doc.add_paragraph(safe_line)
        para.paragraph_format.space_after = Pt(2)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Daily report builder
# ---------------------------------------------------------------------------

def rgb_from_hex(hex_str):
    from docx.shared import RGBColor
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _infer_tag_from_filename_local(fname: str) -> str:
    """Try to infer a district-category tag from the filename."""
    base = os.path.splitext(fname)[0].upper()
    for code in ["TVM", "KLM", "PTA", "ALP", "KTM", "IDK", "EKM",
                  "TSR", "PKD", "MPM", "KKD", "WYD", "KNR", "KSD"]:
        if code in base:
            return f"({code})"
    return ""


def add_fgps_alert_table(doc, report_date: str, failed_files: list):
    """Add a formal government presentation spec table for corrupted file notifications."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 5% grey shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    tcPr.append(shd)
    
    # Dark red 1.5pt border (#C00000)
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # 1.5pt = 12 eighths of pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'C00000')
        borders.append(border)
    tcPr.append(borders)
    
    # Internal margins 0.2cm (113 dxa)
    tcMar = OxmlElement('w:tcMar')
    for margin_name in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), '113')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    run_title = p.add_run("ADMINISTRATIVE NOTE — DATA INTEGRITY ALERT\n")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = rgb_from_hex("C00000")
    
    run_desc = p.add_run(
        f"The consolidation pipeline for {report_date} completed with isolated coverage exceptions. \n"
        "The following district sources were unreadable or malformed and could not be consolidated:\n"
    )
    run_desc.font.name = 'Times New Roman'
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = rgb_from_hex("333333")
    
    for f in failed_files:
        section = f.get("section", "Unknown Section")
        filename = f.get("filename", "Unknown file")
        district = _infer_tag_from_filename_local(filename).strip("()")
        run_file = p.add_run(f"  - {section}: District {district} ({filename})\n")
        run_file.font.name = 'Times New Roman'
        run_file.font.size = Pt(10)
        run_file.font.color.rgb = rgb_from_hex("333333")
        
    run_advise = p.add_run(
        "Officers are advised to perform manual verification for the affected jurisdictions. "
        "All other districts consolidated successfully."
    )
    run_advise.font.name = 'Times New Roman'
    run_advise.font.size = Pt(10)
    run_advise.font.color.rgb = rgb_from_hex("333333")


def validate_report_structure(input_counts: dict, output_counts: dict) -> list:
    """Compare input item counts against output paragraph counts.
    Returns list of warnings if mismatches detected."""
    warnings = []
    for section in ["events", "forecasts", "social_media"]:
        in_c = input_counts.get(section, 0)
        out_c = output_counts.get(section, 0)
        if in_c != out_c:
            warnings.append(f"{section.capitalize()} count mismatch: {in_c} input, {out_c} output")
    return warnings


def verify_generated_report_paragraphs(report_path: str) -> dict:
    """Read the generated report and count the number of items in each section."""
    doc = Document(report_path)
    counts = {"events": 0, "forecasts": 0, "social_media": 0}
    current_section = None
    
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        
        # Identify section headers
        if txt.startswith("II.") or "FORECAST" in txt.upper():
            current_section = "forecasts"
            continue
        elif txt.startswith("III.") or "SOCIAL MEDIA" in txt.upper():
            current_section = "social_media"
            continue
        elif "Important Events/Activities/Issues" in txt:
            current_section = "events"
            continue
            
        # Count items
        if current_section and re.match(r"^\d+[)\t]", txt):
            counts[current_section] += 1
            
    return counts


def build_daily_report(
    report_date: str,
    events: list,
    forecasts: list,
    social_media: list,
    output_path: str,
    ref_number: str = "",
    failed_files: list = None,
):
    """Build a consolidated IS Daily Report .docx.

    Parameters
    ----------
    report_date : str
        Date string in dd.mm.yyyy format.
    events : list of str
        Section I event paragraphs (already in English).
    forecasts : list of str
        Section II forecast paragraphs.
    social_media : list of str
        Section III social media items.
    output_path : str
        Where to save the output .docx.
    ref_number : str
        Optional reference number like "No.961/22/ISDR/SB".
    failed_files : list
        List of failed files for FGPS alerts.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # --- Header ---
    p = doc.add_paragraph("// Secret//")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remove the data integrity alert at the top of the page
    # if failed_files:
    #     add_fgps_alert_table(doc, report_date, failed_files)
    #     doc.add_paragraph()  # blank line after table

    if ref_number:
        p2 = doc.add_paragraph(ref_number)
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # blank line

    p3 = doc.add_paragraph("Daily Report of IS Division, SSB, Kerala Police")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # blank line

    # Compute previous date for the header
    try:
        dt = datetime.strptime(report_date, "%d.%m.%Y")
        prev_dt = dt - timedelta(days=1)
        prev_date_str = prev_dt.strftime("%d.%m.%Y")
    except ValueError:
        prev_date_str = "XX.XX.XXXX"

    header_text = f"Important Events/Activities/Issues ({prev_date_str} 18.00 hrs To {report_date} 18.00hrs)"
    doc.add_paragraph(header_text)

    # --- Section I: Events ---
    for i, event in enumerate(events, 1):
        doc.add_paragraph(f"{i})\t\t{event}")

    # --- Section II: Forecast ---
    if forecasts:
        # Compute a forecast range
        try:
            dt = datetime.strptime(report_date, "%d.%m.%Y")
            next_day = dt + timedelta(days=1)
            end_day = dt + timedelta(days=15)
            forecast_header = f"II. \tFORECAST – ({next_day.strftime('%d.%m.%Y')} – {end_day.strftime('%d.%m.%Y')})"
        except ValueError:
            forecast_header = "II. \tFORECAST"

        doc.add_paragraph()  # blank line
        doc.add_paragraph(forecast_header)

        for i, forecast in enumerate(forecasts, 1):
            doc.add_paragraph(f"{i})\t\t{forecast}")

    # --- Section III: Social Media ---
    if social_media:
        sm_header = f"III.\tSOCIAL MEDIA REPORTS  FROM RSU - ({report_date})"
        doc.add_paragraph()  # blank line
        doc.add_paragraph(sm_header)

        for i, sm_item in enumerate(social_media, 1):
            doc.add_paragraph(f"{i})\t\t{sm_item}")

    # --- Footer ---
    doc.add_paragraph()
    p_footer1 = doc.add_paragraph("Yours faithfully,")
    p_footer1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    p_footer2 = doc.add_paragraph("DySP, DSDB  for SP (IS)")
    p_footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_path)
    return output_path


def extract_person_names(text: str) -> list:
    """Extract potential person names from English report text using heuristics, filtering parent names."""
    candidates = extract_candidate_names_from_text(text)
    
    # Extract parent names S/o, D/o, W/o to ensure we don't treat parent names as new suspect profiles
    parent_names = re.findall(r"[SsDd]/o\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,2})", text)
    parent_names_clean = []
    
    for pn in parent_names:
        words = [w for w in pn.split() if w.lower() not in HONORIFICS]
        if words:
            parent_names_clean.append(" ".join(words))
            
    filtered = []
    for cand in candidates:
        is_parent = False
        for pn in parent_names_clean:
            if is_fuzzy_match(cand, pn):
                is_parent = True
                break
        if is_parent:
            continue

        # Check if the name in text is immediately followed by a police rank/title (cops don't get profiles)
        pattern = r"\b" + re.escape(cand) + r"\b\s*,?\s*\b(ACP|DYSP|SP|SI|CI|Inspector|Commissioner|Superintendent|Deputy\s+Superintendent|Assistant\s+Commissioner|Sub\s+Inspector|Circle\s+Inspector)\b"
        if re.search(pattern, text, re.IGNORECASE):
            continue

        filtered.append(cand)
            
    return filtered


# ---------------------------------------------------------------------------
# Convenience: parse social media file into items
# ---------------------------------------------------------------------------

def parse_social_media_file(path: str) -> list:
    """Parse a SOCIALMEDIA.docx into a list of individual report strings."""
    doc = Document(path)
    items = []
    current_item = ""

    # Try table-based format first
    if doc.tables:
        for table in doc.tables:
            for row in table.rows[1:]:  # skip header row
                cells = [cell.text.strip() for cell in row.cells]
                # Typically: Sl No | Name | Gist | Remarks
                if len(cells) >= 3:
                    name = cells[1] if len(cells) > 1 else ""
                    gist = cells[2] if len(cells) > 2 else ""
                    if gist:
                        item = f"{name}, {gist}" if name else gist
                        items.append(item)
        if items:
            return items

    # Fallback: paragraph-based parsing (numbered items)
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        # Check if this is a new numbered item
        if re.match(r"^\d+[).]?\s", txt):
            if current_item:
                items.append(current_item)
            current_item = re.sub(r"^\d+[).]?\s*", "", txt)
        else:
            current_item += " " + txt

    if current_item:
        items.append(current_item)

    return items


def extract_details_from_docx_paragraphs(paragraphs: list) -> str:
    """Isolate the core narrative paragraphs of a report, discarding
    administrative headers, subjects, and footers.
    """
    # Clean paragraphs
    cleaned_paras = [p.strip() for p in paragraphs if p.strip()]
    
    header_prefixes = (
        "no:", "no.", "നമ്പർ:", "പ്രേഷകൻ:", "സ്വീകർത്താവ്:", "തീയതി:", 
        "date:", "time:", "to:", "from:", "info:", "sir,", "cob message", 
        "ഇമെയിൽ സന്ദേശം", "email message", "special branch", "field report", "ഫീൽഡ് റിപ്പോർട്ട്"
    )
    
    footer_prefixes = (
        "വിശ്വസ്തതയോടെ", "yours faithfully", "yours sincerely", 
        "for dysp", "dysp", "ഡെപ്യൂട്ടി സൂപ്രണ്ട്", "copy to", "പകർപ്പ്", "source:"
    )
    
    details_start_idx = 0
    subject_found = False
    
    # First check for subject line
    for idx, p in enumerate(cleaned_paras):
        p_low = p.lower()
        if p_low.startswith(("വിഷയം:", "sub:", "subject:", "വിശദവിവരങ്ങൾ:", "വിവരം:")) or p_low == "വിഷയം" or p_low == "sub":
            details_start_idx = idx + 1
            subject_found = True
            break
    
    # If no subject was found, skip leading headers
    if not subject_found:
        for idx, p in enumerate(cleaned_paras):
            p_low = p.lower()
            is_header = False
            if p_low.startswith(header_prefixes):
                is_header = True
            elif "internal security input report" in p_low:
                is_header = True
            elif "alert report" in p_low:
                is_header = True
            
            if is_header:
                details_start_idx = idx + 1
            else:
                break
    
    # Check if there is an "Elaborated Report" / "Detailed Description" after the start index
    for idx in range(details_start_idx, len(cleaned_paras)):
        p_low = cleaned_paras[idx].lower()
        if "elaborated report" in p_low or "detailed description" in p_low or "detailed narrative" in p_low:
            details_start_idx = idx + 1
            break
            
    # Gather paragraphs from details_start_idx to the footer
    details_paras = []
    for idx in range(details_start_idx, len(cleaned_paras)):
        p = cleaned_paras[idx]
        p_low = p.lower()
        # Check for footer
        if any(p_low.startswith(f) for f in footer_prefixes) or p_low == "വിശ്വസ്തതയോടെ,":
            break
        details_paras.append(p)
        
    return "\n".join(details_paras)


def build_less_priority_report(
    report_date: str,
    items: list,
    output_path: str,
):
    """Build a consolidated DOCX file for not important / less priority reports.

    Parameters
    ----------
    report_date : str
        Date string in dd.mm.yyyy format.
    items : list of str
        Unimportant event paragraphs (in English).
    output_path : str
        Where to save the output .docx.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # --- Header ---
    p = doc.add_paragraph("// Secret//")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # blank line

    p3 = doc.add_paragraph("Daily Less Priority / Not Needed Reports for Human Verification")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # blank line

    header_text = f"Unimportant / Not Needed Reports ({report_date})"
    doc.add_paragraph(header_text)

    # --- Items ---
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i})\t\t{item}")

    # --- Footer ---
    doc.add_paragraph()
    p_footer1 = doc.add_paragraph("Yours faithfully,")
    p_footer1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    p_footer2 = doc.add_paragraph("DySP, DSDB  for SP (IS)")
    p_footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_path)
    return output_path

